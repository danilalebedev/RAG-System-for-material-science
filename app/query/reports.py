from __future__ import annotations

import csv
import hashlib
import json
import os
import re
import shutil
import subprocess
import zipfile
from pathlib import Path
from types import SimpleNamespace
from typing import Any
from xml.sax.saxutils import escape

from docx import Document
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


PROJECT_ROOT = Path(__file__).resolve().parents[2]
MOJIBAKE_MARKERS = ("Р", "С", "Ð", "Ñ", "в†", "В°", "Вµ")
MAX_LOCAL_ARCHIVE_FILES = 20
MAX_LOCAL_ARCHIVE_BYTES = 250 * 1024 * 1024
DEFAULT_ROUTERAI_BUDGET_RUB = 1500.0
XML_INVALID_CHAR_RE = re.compile(r"[\x00-\x08\x0B\x0C\x0E-\x1F\uD800-\uDFFF]")
DEFAULT_ROUTERAI_PROMPT_RUB_PER_1K = 0.03
DEFAULT_ROUTERAI_COMPLETION_RUB_PER_1K = 0.12


def repair_mojibake(value: Any) -> str:
    text = str(value or "")
    if not text or not any(marker in text for marker in MOJIBAKE_MARKERS):
        return text
    try:
        fixed = text.encode("cp1251").decode("utf-8")
    except UnicodeError:
        return text
    if fixed.count("�") > text.count("�"):
        return text
    return fixed


def compact_text(value: Any, max_chars: int | None = None) -> str:
    if isinstance(value, (dict, list)):
        text = json.dumps(value, ensure_ascii=False, default=str)
    else:
        text = repair_mojibake(value)
    text = XML_INVALID_CHAR_RE.sub(" ", text)
    text = re.sub(r"\s+", " ", text).strip()
    if max_chars is not None and len(text) > max_chars:
        return text[: max_chars - 3].rstrip() + "..."
    return text


def list_values(value: Any, *, max_chars: int = 180) -> list[str]:
    if value in (None, "", [], {}):
        return []
    values = value if isinstance(value, list) else [value]
    result: list[str] = []
    for item in values:
        text = compact_text(item, max_chars)
        if text:
            result.append(text)
    return result


def unique_limited(values: list[str], limit: int = 8) -> list[str]:
    result: list[str] = []
    seen: set[str] = set()
    for value in values:
        text = compact_text(value, 220)
        key = text.casefold()
        if not text or key in seen:
            continue
        seen.add(key)
        result.append(text)
        if len(result) >= limit:
            break
    return result


def result_link(result: Any) -> str:
    if getattr(result, "url", None):
        return str(result.url)
    if getattr(result, "doi", None):
        doi = str(result.doi).removeprefix("https://doi.org/")
        return f"https://doi.org/{doi}"
    return ""


def display_query(run: Any) -> str:
    plan = getattr(run, "query_plan", {}) or {}
    rewrite = plan.get("llm_rewrite") if isinstance(plan.get("llm_rewrite"), dict) else {}
    return compact_text(
        rewrite.get("corrected_query")
        or plan.get("corrected_query")
        or plan.get("original_query")
        or getattr(getattr(run, "request", None), "query", "")
    )


def source_title(row: dict[str, Any]) -> str:
    return compact_text(row.get("title") or row.get("doc_id") or row.get("source_path") or row.get("local_path") or "Источник", 300)


def source_link(result: Any) -> str:
    link = result_link(result)
    return link or compact_text(getattr(result, "doi", "") or getattr(result, "result_id", ""))


def preferred_web_link(result: Any) -> str:
    return compact_text(
        getattr(result, "open_access_pdf_url", None)
        or (getattr(result, "open_access", {}) or {}).get("best_pdf_url")
        or source_link(result),
        1000,
    )


def llm_comparison_summary(run: Any) -> str:
    plan = getattr(run, "query_plan", {}) or {}
    return compact_text(plan.get("llm_comparison_summary"), 6000)


def result_quartile(result: Any) -> str:
    raw = getattr(result, "raw", None) or {}
    return compact_text(getattr(result, "journal_quartile", None) or raw.get("journal_quartile"), 20)


def relevance_confidence(result: Any) -> dict[str, Any]:
    raw = getattr(result, "raw", None) or {}
    score = float(getattr(result, "score", 0.0) or 0.0)
    keyword_hits = list(getattr(result, "keyword_hits", []) or [])
    title_hits = list(raw.get("title_relevance_hits") or [])
    snippet_hits = list(raw.get("snippet_relevance_hits") or [])
    domain_hits = list(raw.get("materials_domain_hits") or [])
    quartile = result_quartile(result)

    percent = min(98, max(5, int(round(score * 3.2))))
    strong_quartile = quartile in {"Q1", "Q2"}
    if score >= 22 or (score >= 16 and (strong_quartile or len(keyword_hits) >= 3)):
        label = "Высокая"
        percent = max(percent, 72)
    elif score >= 10 or keyword_hits or domain_hits:
        label = "Средняя"
        percent = max(percent, 42)
    else:
        label = "Низкая"
        percent = min(percent, 39)

    reasons: list[str] = []
    if keyword_hits:
        reasons.append("ключевые совпадения: " + ", ".join(compact_text(item, 60) for item in keyword_hits[:5]))
    if title_hits:
        reasons.append("совпадения в заголовке: " + ", ".join(compact_text(item, 60) for item in title_hits[:4]))
    if snippet_hits:
        reasons.append("совпадения в abstract/snippet: " + ", ".join(compact_text(item, 60) for item in snippet_hits[:4]))
    if domain_hits:
        reasons.append("materials-domain signals: " + ", ".join(compact_text(item, 60) for item in domain_hits[:4]))
    if quartile:
        reasons.append(f"квартиль журнала: {quartile}")
    if getattr(result, "doi", None):
        reasons.append("есть DOI")
    if getattr(result, "abstract", None):
        reasons.append("есть abstract")
    if getattr(result, "citation_count", None):
        reasons.append(f"цитирований: {getattr(result, 'citation_count')}")
    if not reasons:
        reasons.append("мало проверяемых metadata-сигналов; нужна ручная проверка")

    return {
        "label": label,
        "confidence": percent,
        "score": round(score, 3),
        "reasons": unique_limited(reasons, limit=6),
    }


def relevance_confidence_text(result: Any) -> str:
    confidence = relevance_confidence(result)
    reasons = "; ".join(confidence["reasons"])
    return f"{confidence['label']} ({confidence['confidence']}%); score={confidence['score']}; {reasons}"


def year_counts(run: Any) -> list[dict[str, Any]]:
    counts: dict[int, int] = {}
    for result in getattr(run, "results", []) or []:
        if getattr(result, "year", None):
            counts[int(result.year)] = counts.get(int(result.year), 0) + 1
    return [{"year": year, "count": counts[year]} for year in sorted(counts)]


def source_counts(run: Any) -> list[dict[str, Any]]:
    counts: dict[str, int] = {}
    for result in getattr(run, "results", []) or []:
        source = compact_text(getattr(result, "source", "unknown"), 80) or "unknown"
        counts[source] = counts.get(source, 0) + 1
    return [{"source": source, "count": count} for source, count in sorted(counts.items())]


def run_overall_summary(run: Any) -> str:
    deep_results = list(getattr(run, "deep_results", []) or [])
    if not deep_results:
        return (
            "Выполнен metadata-only поиск: сформирован ранжированный список источников и ссылок. "
            "Запустите Deep Search, чтобы извлечь краткие summary статей и сравнить методики с локальной базой."
        )

    summaries = [item.document_summary for item in deep_results if getattr(item, "document_summary", None)]
    procedures_count = sum(len(getattr(item, "procedure_summaries", []) or []) for item in deep_results)
    comparison = getattr(run, "comparison", None)
    confirmed = len(getattr(comparison, "confirmed_methods", []) or []) if comparison else 0
    web_only = len(getattr(comparison, "web_only_methods", []) or []) if comparison else 0
    local_only = len(getattr(comparison, "local_only_methods", []) or []) if comparison else 0

    paper_summaries = unique_limited(
        [compact_text(row.get("summary") or row.get("main_topic"), 320) for row in summaries if row],
        limit=4,
    )
    materials = unique_limited([item for row in summaries for item in list_values(row.get("materials"))], limit=8)
    processes = unique_limited(
        [item for row in summaries for item in list_values(row.get("processes") or row.get("methods"))],
        limit=8,
    )
    findings = unique_limited([item for row in summaries for item in list_values(row.get("key_findings"))], limit=5)

    parts = [
        f"Deep Search обработал {len(summaries)} внешних источников и извлек {procedures_count} записей о методиках."
    ]
    if paper_summaries:
        parts.append("Общий вывод по статьям: " + " ".join(paper_summaries))
    if materials:
        parts.append("Основные материалы: " + ", ".join(materials) + ".")
    if processes:
        parts.append("Основные процессы/методики: " + ", ".join(processes) + ".")
    if findings:
        parts.append("Ключевые наблюдения: " + "; ".join(findings) + ".")
    if comparison:
        parts.append(
            "Сравнение с локальной базой: "
            f"подтверждено методик {confirmed}, найдено только во внешней литературе {web_only}, "
            f"только локально {local_only}."
        )
    return " ".join(parts)


def comparison_insights(run: Any) -> str:
    request = getattr(run, "request", None)
    if request is not None and not getattr(request, "generate_comparison_insights", True):
        return ""
    llm_summary = llm_comparison_summary(run)
    if llm_summary:
        return llm_summary
    comparison = getattr(run, "comparison", None)
    if comparison is None:
        return "Сравнение локального и web-поиска пока недоступно: нет comparison report."

    years = year_counts(run)
    sources = source_counts(run)
    newest = max((row["year"] for row in years), default=None)
    oldest = min((row["year"] for row in years), default=None)
    top_source = max(sources, key=lambda row: row["count"], default=None)
    confirmed = len(getattr(comparison, "confirmed_methods", []) or [])
    local_only = len(getattr(comparison, "local_only_methods", []) or [])
    web_only = len(getattr(comparison, "web_only_methods", []) or [])
    differing = len(getattr(comparison, "differing_conditions", []) or [])

    parts: list[str] = []
    if oldest and newest:
        parts.append(f"Внешняя выдача покрывает публикации за период {oldest}-{newest}.")
    if top_source:
        parts.append(f"Больше всего публикаций пришло из базы {top_source['source']} ({top_source['count']}).")
    parts.append(
        f"По методикам: подтверждено пересечений {confirmed}, локально уникальных записей {local_only}, "
        f"внешне уникальных записей {web_only}."
    )
    if differing:
        parts.append(f"Найдено {differing} случаев, где методики похожи, но условия или численные диапазоны отличаются.")
    if web_only > local_only:
        parts.append("Внешняя литература расширяет локальную базу и может дать кандидатов для пополнения графа знаний.")
    elif local_only > web_only:
        parts.append("Локальная база содержит больше уникальных методик по запросу; web-поиск полезен для подтверждения и бенчмарка.")
    else:
        parts.append("Локальная и внешняя выдача дают сопоставимый объем уникальных методик.")
    return " ".join(parts)


def build_links_report(run: Any) -> str:
    lines = [
        "# Отчет по релевантным ссылкам",
        "",
        f"Запрос: {compact_text(run.request.query)}",
        "",
        "## Web-search",
    ]
    if getattr(run, "results", None):
        for index, result in enumerate(run.results, start=1):
            title = compact_text(result.title, 300)
            link = source_link(result)
            lines.append(f"{index}. [{title}]({link})" if link.startswith("http") else f"{index}. {title} - {link}")
            lines.append(f"   - Confidence: {relevance_confidence_text(result)}")
    else:
        lines.append("Внешние публикации не найдены.")

    lines.extend(["", "## Локальный поиск"])
    local_matches = list(getattr(run, "local_matches", []) or [])
    if local_matches:
        for index, row in enumerate(local_matches[:50], start=1):
            lines.append(f"{index}. {source_title(row)}")
    else:
        lines.append("Локальные совпадения не найдены.")
    return "\n".join(lines).strip() + "\n"


def build_deep_report(run: Any) -> str:
    lines = [
        "# Deep Search отчет",
        "",
        f"Запрос: {compact_text(run.request.query)}",
        "",
        "## Общий вывод",
        "",
        run_overall_summary(run),
        "",
        "## Summary по статьям",
    ]
    deep_results = list(getattr(run, "deep_results", []) or [])
    if not deep_results:
        lines.append("Deep Search еще не запускался.")
    for index, item in enumerate(deep_results, start=1):
        summary = item.document_summary or {}
        link = source_link(item.source_result)
        text = compact_text(summary.get("summary") or summary.get("main_topic") or "Summary не извлечен.", 2000)
        lines.extend(["", f"### {index}. {compact_text(item.source_result.title, 260)}", f"Ссылка: {link or 'n/a'}", text])
    return "\n".join(lines).strip() + "\n"


def build_executive_brief_report(run: Any) -> str:
    comparison = getattr(run, "comparison", None)
    findings = [
        f"Web-источников: {len(getattr(run, 'results', []) or [])}.",
        f"Локальных совпадений: {len(getattr(run, 'local_matches', []) or [])}.",
        f"Deep Search summary: {len(getattr(run, 'deep_results', []) or [])}.",
    ]
    if comparison:
        findings.append(
            f"Методики: confirmed={len(comparison.confirmed_methods)}, local-only={len(comparison.local_only_methods)}, "
            f"web-only={len(comparison.web_only_methods)}."
        )
    lines = ["# Краткий управленческий вывод", "", f"Запрос: {compact_text(run.request.query)}", "", "## 5 ключевых выводов"]
    lines.extend(f"- {item}" for item in findings[:5])
    if getattr(run, "results", None):
        lines.extend(["", "## Самые релевантные ссылки"])
        for index, result in enumerate(run.results[:8], start=1):
            lines.append(f"{index}. {compact_text(result.title, 220)} - {source_link(result)}")
            lines.append(f"   - Confidence: {relevance_confidence_text(result)}")
    return "\n".join(lines).strip() + "\n"


def build_literature_report(run: Any) -> str:
    lines = [
        "# Полный отчет по поиску литературы",
        "",
        f"Запрос: {compact_text(run.request.query)}",
        f"Ключевые слова: {', '.join(compact_text(item, 80) for item in getattr(run, 'keywords', []) or []) or 'n/a'}",
        f"Внешние результаты: {len(getattr(run, 'results', []) or [])}",
        f"Локальные совпадения: {len(getattr(run, 'local_matches', []) or [])}",
        f"Deep Search summaries: {len(getattr(run, 'deep_results', []) or [])}",
        "",
        "## Общий вывод",
        "",
        run_overall_summary(run),
    ]
    insights = comparison_insights(run)
    if insights:
        lines.extend(["", "## Выводы по сравнению локального и web-поиска", "", insights])

    lines.extend(["", "## Локальный поиск"])
    local_matches = list(getattr(run, "local_matches", []) or [])
    if local_matches:
        for index, row in enumerate(local_matches[:50], start=1):
            lines.append(f"{index}. {source_title(row)}")
    else:
        lines.append("Локальные совпадения не найдены.")

    lines.extend(["", "## Web-search"])
    if getattr(run, "results", None):
        for index, result in enumerate(run.results, start=1):
            lines.append(f"{index}. [{compact_text(result.title, 300)}]({source_link(result)})")
            lines.append(f"   - Confidence: {relevance_confidence_text(result)}")
    else:
        lines.append("Внешние публикации не найдены.")

    lines.extend(["", "## Графики и распределения", "", "### Публикации по годам"])
    year_rows = year_counts(run)
    lines.extend([f"- {row['year']}: {row['count']}" for row in year_rows] or ["- Нет данных по годам."])
    lines.extend(["", "### Публикации по базам данных"])
    source_rows = source_counts(run)
    lines.extend([f"- {row['source']}: {row['count']}" for row in source_rows] or ["- Нет данных по источникам."])

    if getattr(run, "deep_results", None):
        lines.extend(["", build_deep_report(run).strip()])
    warnings = list(getattr(run, "warnings", []) or [])
    if warnings:
        lines.extend(["", "## Warnings"])
        lines.extend(f"- {compact_text(warning, 500)}" for warning in warnings)
    return "\n".join(lines).strip() + "\n"


def build_full_run_payload(run: Any) -> dict[str, Any]:
    return {
        "request": run.request.model_dump(mode="json"),
        "query_plan": getattr(run, "query_plan", {}) or {},
        "keywords": list(getattr(run, "keywords", []) or []),
        "web_results": [row.model_dump(mode="json") for row in getattr(run, "results", []) or []],
        "local_matches": list(getattr(run, "local_matches", []) or []),
        "deep_results": [row.model_dump(mode="json") for row in getattr(run, "deep_results", []) or []],
        "comparison": run.comparison.model_dump(mode="json") if getattr(run, "comparison", None) else None,
        "charts": {
            "publication_years": year_counts(run),
            "sources": source_counts(run),
            "local_vs_web": [
                {"bucket": "local", "count": len(getattr(run, "local_matches", []) or [])},
                {"bucket": "web", "count": len(getattr(run, "results", []) or [])},
            ],
        },
        "overall_summary": run_overall_summary(run),
        "comparison_insights": comparison_insights(run),
        "warnings": list(getattr(run, "warnings", []) or []),
    }


def register_pdf_font() -> str:
    candidates = [
        Path("C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibri.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
    ]
    for path in candidates:
        if path.exists():
            font_name = path.stem.replace(" ", "_")
            if font_name not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(TTFont(font_name, str(path)))
            return font_name
    return "Helvetica"


def pdf_styles() -> dict[str, ParagraphStyle]:
    font_name = register_pdf_font()
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="BodyUnicode", parent=styles["BodyText"], fontName=font_name, fontSize=9, leading=12))
    styles.add(ParagraphStyle(name="TitleUnicode", parent=styles["Title"], fontName=font_name, fontSize=16, leading=20))
    styles.add(ParagraphStyle(name="HeadingUnicode", parent=styles["Heading2"], fontName=font_name, fontSize=12, leading=15))
    styles.add(ParagraphStyle(name="SmallUnicode", parent=styles["BodyText"], fontName=font_name, fontSize=8, leading=10))
    return styles


def paragraph(text: Any, style: ParagraphStyle) -> Paragraph:
    return Paragraph(escape(compact_text(text, 4000)) or " ", style)


def link_paragraph(label: str, url: str | None, style: ParagraphStyle) -> Paragraph:
    safe_label = escape(compact_text(label, 400))
    safe_url = escape(str(url or ""))
    if safe_url.startswith(("http://", "https://")):
        return Paragraph(f'<link href="{safe_url}">{safe_label}</link>', style)
    return Paragraph(safe_label, style)


def add_sources_table(story: list[Any], run: Any, styles: dict[str, ParagraphStyle]) -> None:
    story.append(paragraph("Релевантные источники", styles["HeadingUnicode"]))
    table_rows = [["#", "Заголовок", "Ссылка"]]
    for index, result in enumerate((getattr(run, "results", []) or [])[:40], start=1):
        link = source_link(result)
        table_rows.append(
            [
                str(index),
                paragraph(result.title, styles["SmallUnicode"]),
                link_paragraph("Открыть", link, styles["SmallUnicode"]) if link.startswith("http") else paragraph(link, styles["SmallUnicode"]),
            ]
        )
    if len(table_rows) == 1:
        table_rows.append(["-", paragraph("Внешние публикации не найдены", styles["SmallUnicode"]), ""])
    table = Table(table_rows, colWidths=[0.8 * cm, 11.4 * cm, 4.1 * cm], repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e8eef7")),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#c8cdd6")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    story.append(table)


def add_count_table(story: list[Any], title: str, headers: list[str], rows: list[dict[str, Any]], styles: dict[str, ParagraphStyle]) -> None:
    if not rows:
        return
    story.append(Spacer(1, 0.25 * cm))
    story.append(paragraph(title, styles["HeadingUnicode"]))
    table_rows = [headers] + [[compact_text(value, 160) for value in row.values()] for row in rows]
    table = Table(table_rows, colWidths=[8 * cm, 3 * cm], repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e8eef7")),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#c8cdd6")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    story.append(table)


def add_pdf_table(story: list[Any], headers: list[str], rows: list[list[Any]], styles: dict[str, ParagraphStyle]) -> None:
    if not rows:
        return
    table_rows: list[list[Any]] = [[paragraph(header, styles["SmallUnicode"]) for header in headers]]
    for row in rows:
        table_rows.append([paragraph(value, styles["SmallUnicode"]) for value in row[: len(headers)]])
    if len(headers) == 5:
        col_widths = [0.7 * cm, 6.9 * cm, 1.3 * cm, 2.0 * cm, 5.2 * cm]
    elif len(headers) == 3:
        col_widths = [0.7 * cm, 6.5 * cm, 8.9 * cm]
    else:
        col_widths = [16.1 * cm / max(len(headers), 1)] * len(headers)
    table = Table(table_rows, colWidths=col_widths, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e8eef7")),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#c8cdd6")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    story.append(table)


def build_pdf_report(run: Any, output_path: Path, *, mode: str = "full") -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    styles = pdf_styles()
    story: list[Any] = []
    titles = {
        "full": "Полный отчет по поиску литературы",
        "links": "Отчет по релевантным ссылкам",
        "deep": "Deep Search отчет",
        "brief": "Краткий управленческий вывод",
    }
    story.append(paragraph(titles.get(mode, "Отчет по поиску литературы"), styles["TitleUnicode"]))
    story.append(paragraph(f"Запрос: {run.request.query}", styles["BodyUnicode"]))

    if mode in {"full", "brief", "deep"}:
        story.append(Spacer(1, 0.25 * cm))
        story.append(paragraph("Общий вывод", styles["HeadingUnicode"]))
        story.append(paragraph(run_overall_summary(run), styles["BodyUnicode"]))
    if mode == "full":
        insights = comparison_insights(run)
        if insights:
            story.append(paragraph("Сравнение локального и web-поиска", styles["HeadingUnicode"]))
            story.append(paragraph(insights, styles["BodyUnicode"]))
    if mode in {"full", "links", "brief"}:
        story.append(Spacer(1, 0.25 * cm))
        add_sources_table(story, run, styles)
    if mode == "full":
        add_count_table(story, "Публикации по годам", ["Год", "Количество"], year_counts(run), styles)
        add_count_table(story, "Публикации по базам данных", ["База", "Количество"], source_counts(run), styles)
    if mode in {"full", "deep"} and getattr(run, "deep_results", None):
        story.append(Spacer(1, 0.4 * cm))
        story.append(paragraph("Deep Search summaries", styles["HeadingUnicode"]))
        for item in run.deep_results:
            summary = item.document_summary or {}
            story.append(link_paragraph(item.source_result.title, source_link(item.source_result), styles["BodyUnicode"]))
            story.append(paragraph(summary.get("summary") or summary.get("main_topic") or "Summary не извлечен.", styles["SmallUnicode"]))

    doc = SimpleDocTemplate(str(output_path), pagesize=A4, rightMargin=1.2 * cm, leftMargin=1.2 * cm, topMargin=1.2 * cm, bottomMargin=1.2 * cm)
    doc.build(story)
    return output_path


def set_docx_style(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)


def add_docx_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = compact_text(value, 1200)


def build_docx_report(run: Any, output_path: Path, *, mode: str = "full") -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    set_docx_style(document)
    titles = {
        "full": "Полный отчет по поиску литературы",
        "links": "Отчет по релевантным ссылкам",
        "deep": "Deep Search отчет",
        "brief": "Краткий управленческий вывод",
    }
    document.add_heading(titles.get(mode, "Отчет по поиску литературы"), level=1)
    document.add_paragraph(f"Запрос: {compact_text(run.request.query)}")

    if mode in {"full", "brief", "deep"}:
        document.add_heading("Общий вывод", level=2)
        document.add_paragraph(compact_text(run_overall_summary(run), 4000))
    if mode == "full":
        insights = comparison_insights(run)
        if insights:
            document.add_heading("Сравнение локального и web-поиска", level=2)
            document.add_paragraph(compact_text(insights, 4000))
    if mode in {"full", "links", "brief"}:
        document.add_heading("Релевантные ссылки", level=2)
        add_docx_table(
            document,
            ["#", "Заголовок", "Ссылка"],
            [[str(index), result.title, source_link(result)] for index, result in enumerate(getattr(run, "results", []) or [], start=1)],
        )
    if mode == "full":
        document.add_heading("Локальный поиск", level=2)
        add_docx_table(
            document,
            ["#", "Источник"],
            [[str(index), source_title(row)] for index, row in enumerate((getattr(run, "local_matches", []) or [])[:50], start=1)],
        )
        document.add_heading("Публикации по годам", level=2)
        add_docx_table(document, ["Год", "Количество"], [[str(row["year"]), str(row["count"])] for row in year_counts(run)])
        document.add_heading("Публикации по базам данных", level=2)
        add_docx_table(document, ["База", "Количество"], [[str(row["source"]), str(row["count"])] for row in source_counts(run)])
    if mode in {"full", "deep"} and getattr(run, "deep_results", None):
        document.add_heading("Deep Search summaries", level=2)
        for index, item in enumerate(run.deep_results, start=1):
            summary = item.document_summary or {}
            document.add_heading(f"{index}. {compact_text(item.source_result.title, 180)}", level=3)
            document.add_paragraph(f"Ссылка: {compact_text(source_link(item.source_result) or 'n/a')}")
            document.add_paragraph(compact_text(summary.get("summary") or summary.get("main_topic") or "Summary не извлечен.", 2200))

    document.save(output_path)
    return output_path


def write_text_report(path: Path, text: str) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding="utf-8")
    return path


def write_json_report(path: Path, payload: dict[str, Any]) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2, default=str), encoding="utf-8")
    return path


def safe_report_id(value: Any, *, prefix: str = "run") -> str:
    text = compact_text(value, 220) or prefix
    digest = hashlib.sha256(text.encode("utf-8", errors="ignore")).hexdigest()[:10]
    slug = re.sub(r"[^A-Za-z0-9А-Яа-яЁё_.-]+", "_", text).strip("_")
    slug = slug[:50].strip("_") or prefix
    return f"{slug}_{digest}"


def markdown_to_docx(text: str, output_path: Path) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    set_docx_style(document)
    for raw_line in text.splitlines():
        line = compact_text(raw_line.strip(), 4000)
        if not line:
            continue
        if line.startswith("### "):
            document.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            document.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            document.add_heading(line[2:], level=1)
        elif line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s", line):
            document.add_paragraph(line, style="List Number")
        else:
            document.add_paragraph(line)
    document.save(output_path)
    return output_path


def markdown_to_pdf(text: str, output_path: Path) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    styles = pdf_styles()
    story: list[Any] = []
    for raw_line in text.splitlines():
        line = compact_text(raw_line.strip(), 4000)
        if not line:
            story.append(Spacer(1, 0.15 * cm))
            continue
        if line.startswith("# "):
            story.append(paragraph(line[2:], styles["TitleUnicode"]))
        elif line.startswith(("## ", "### ")):
            story.append(paragraph(line.lstrip("# "), styles["HeadingUnicode"]))
        else:
            story.append(paragraph(line, styles["BodyUnicode"]))
    doc = SimpleDocTemplate(str(output_path), pagesize=A4, rightMargin=1.2 * cm, leftMargin=1.2 * cm, topMargin=1.2 * cm, bottomMargin=1.2 * cm)
    doc.build(story)
    return output_path


def literature_graph_markdown(run: Any) -> str:
    from app.query.cockpit import mini_graph_edges

    lines = ["# Граф evidence", "", f"Запрос: {compact_text(run.request.query)}", ""]
    edges = mini_graph_edges(run)
    if not edges:
        lines.append("Graph edges не найдены.")
        return "\n".join(lines).strip() + "\n"
    lines.extend(["## Ребра"])
    for index, edge in enumerate(edges[:80], start=1):
        lines.append(
            f"{index}. {compact_text(edge.get('from'), 180)} --[{compact_text(edge.get('relation'), 80)} / {compact_text(edge.get('scope'), 40)}]--> {compact_text(edge.get('to'), 180)}"
        )
    return "\n".join(lines).strip() + "\n"


def property_report_rows_from_run(run: Any) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    comparison = getattr(run, "comparison", None)
    for row in getattr(comparison, "rows", []) or []:
        rows.append(
            {
                "scope": row.get("scope"),
                "material": row.get("material") or row.get("material_name"),
                "method": row.get("method") or row.get("synthesis_or_process_method"),
                "properties": row.get("properties") or row.get("outputs") or row.get("observed_effects"),
                "numeric_results": row.get("numeric_results") or row.get("numerical_results") or row.get("analysis_results"),
                "conditions": row.get("conditions"),
                "evidence": row.get("title") or row.get("doc_id") or row.get("result_id"),
            }
        )
    for deep_result in getattr(run, "deep_results", []) or []:
        source = getattr(deep_result, "source_result", None)
        for procedure in getattr(deep_result, "procedure_summaries", []) or []:
            rows.append(
                {
                    "scope": "web_deep_search",
                    "material": procedure.get("material_name") or procedure.get("materials"),
                    "method": procedure.get("synthesis_or_process_method") or procedure.get("method"),
                    "properties": procedure.get("properties") or procedure.get("outputs") or procedure.get("observed_effects"),
                    "numeric_results": procedure.get("numeric_results") or procedure.get("numerical_results") or procedure.get("analysis_results"),
                    "conditions": procedure.get("conditions"),
                    "evidence": getattr(source, "title", "") if source is not None else procedure.get("result_id"),
                }
            )
    for row in getattr(run, "local_matches", []) or []:
        if any(row.get(key) for key in ("outputs", "properties", "numeric_results", "analysis_results", "conditions")):
            rows.append(
                {
                    "scope": "local",
                    "material": row.get("material") or row.get("material_name"),
                    "method": row.get("method") or row.get("synthesis_or_process_method"),
                    "properties": row.get("properties") or row.get("outputs") or row.get("observed_effects"),
                    "numeric_results": row.get("numeric_results") or row.get("numerical_results") or row.get("analysis_results"),
                    "conditions": row.get("conditions"),
                    "evidence": row.get("title") or row.get("doc_id") or row.get("source_path"),
                }
            )
    return rows[:120]


def property_report_rows_from_orchestration(orchestration: Any) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    context = orchestration_context(orchestration)
    for section, section_rows in context.items():
        for row in section_rows or []:
            if row.get("source_type") == "diagnostics":
                continue
            payload = row.get("row") if isinstance(row.get("row"), dict) else row
            has_property_signal = any(
                payload.get(key) or row.get(key)
                for key in (
                    "properties",
                    "outputs",
                    "observed_effects",
                    "numeric_results",
                    "numerical_results",
                    "analysis_results",
                    "conditions",
                    "rows",
                    "summary",
                    "preview",
                )
            )
            if not has_property_signal:
                continue
            rows.append(
                {
                    "scope": section,
                    "material": payload.get("material") or payload.get("material_name") or row.get("matched_terms"),
                    "method": payload.get("method") or payload.get("synthesis_or_process_method"),
                    "properties": payload.get("properties") or payload.get("outputs") or payload.get("observed_effects") or row.get("summary"),
                    "numeric_results": payload.get("numeric_results")
                    or payload.get("numerical_results")
                    or payload.get("analysis_results")
                    or row.get("preview")
                    or row.get("rows"),
                    "conditions": payload.get("conditions") or row.get("path"),
                    "evidence": row_title(row),
                }
            )
    return rows[:120]


def property_report_markdown(title: str, query: str, rows: list[dict[str, Any]]) -> str:
    lines = [f"# {title}", "", f"Запрос: {compact_text(query)}", ""]
    if not rows:
        lines.append("Свойства и численные результаты не найдены.")
        return "\n".join(lines).strip() + "\n"
    lines.append("## Свойства, условия и численные результаты")
    for index, row in enumerate(rows, start=1):
        lines.append(f"{index}. {compact_text(row.get('material') or 'material n/a', 180)}; {compact_text(row.get('method') or 'method n/a', 180)}")
        lines.append(f"   - scope: {compact_text(row.get('scope'), 80)}")
        lines.append(f"   - properties/outputs: {compact_text(row.get('properties'), 500)}")
        lines.append(f"   - numeric results: {compact_text(row.get('numeric_results'), 500)}")
        lines.append(f"   - conditions: {compact_text(row.get('conditions'), 500)}")
        lines.append(f"   - evidence: {compact_text(row.get('evidence'), 260)}")
    return "\n".join(lines).strip() + "\n"


def build_section_markdown(run: Any, section: str) -> str:
    section = section.lower().strip()
    if section == "sources":
        return build_links_report(run)
    if section == "deep":
        return build_deep_report(run)
    if section in {"graphs", "graph"}:
        return literature_graph_markdown(run)
    if section in {"properties", "property", "numeric"}:
        return property_report_markdown("Свойства и численные результаты", run.request.query, property_report_rows_from_run(run))
    if section == "charts":
        lines = ["# Распределения публикаций", "", f"Запрос: {compact_text(run.request.query)}", "", "## По годам"]
        lines.extend([f"- {row['year']}: {row['count']}" for row in year_counts(run)] or ["- Нет данных по годам."])
        lines.extend(["", "## По базам данных"])
        lines.extend([f"- {row['source']}: {row['count']}" for row in source_counts(run)] or ["- Нет данных по базам."])
        lines.extend(["", "## Local vs web"])
        lines.append(f"- Local matches: {len(getattr(run, 'local_matches', []) or [])}")
        lines.append(f"- Web results: {len(getattr(run, 'results', []) or [])}")
        lines.append(f"- Deep Search summaries: {len(getattr(run, 'deep_results', []) or [])}")
        return "\n".join(lines).strip() + "\n"
    if section == "comparison":
        comparison = getattr(run, "comparison", None)
        lines = ["# Сравнение локального и web-поиска", "", f"Запрос: {compact_text(run.request.query)}", "", comparison_insights(run)]
        if not comparison:
            lines.append("Comparison report пока недоступен.")
            return "\n".join(lines).strip() + "\n"
        buckets = [
            ("Подтверждается локально и внешне", comparison.confirmed_methods),
            ("Только локально", comparison.local_only_methods),
            ("Только во внешней литературе", comparison.web_only_methods),
            ("Разные условия или диапазоны", comparison.differing_conditions),
        ]
        for title, rows in buckets:
            lines.extend(["", f"## {title}"])
            if not rows:
                lines.append("- Нет данных.")
                continue
            for index, row in enumerate(rows[:30], start=1):
                material = compact_text(row.get("material") or row.get("title") or row.get("local_title") or row.get("web_title"), 180)
                method = compact_text(row.get("method") or row.get("processes"), 180)
                lines.append(f"{index}. {material}; {method}")
        return "\n".join(lines).strip() + "\n"
    if section == "evidence":
        lines = ["# Evidence", "", f"Запрос: {compact_text(run.request.query)}", "", "## Локальные совпадения"]
        for index, row in enumerate((getattr(run, "local_matches", []) or [])[:40], start=1):
            lines.append(f"{index}. {source_title(row)} - {compact_text(row.get('preview') or row.get('source_path'), 300)}")
        if len(lines) == 5:
            lines.append("Локальные совпадения не найдены.")
        lines.extend(["", "## Web evidence"])
        for index, result in enumerate((getattr(run, "results", []) or [])[:40], start=1):
            lines.append(f"{index}. {compact_text(result.title, 220)} - {source_link(result)}")
        return "\n".join(lines).strip() + "\n"
    return build_literature_report(run)


def build_section_exports(run: Any, section: str, output_dir: Path | None = None) -> dict[str, Path]:
    output_dir = output_dir or Path(getattr(run, "output_dir", "") or PROJECT_ROOT / "data" / "processed" / "web_search" / "section_exports")
    output_dir.mkdir(parents=True, exist_ok=True)
    safe_section = re.sub(r"[^A-Za-z0-9_.-]+", "_", section.lower()).strip("_") or "section"
    markdown = build_section_markdown(run, safe_section)
    md_path = write_text_report(output_dir / f"{safe_section}_report.md", markdown)
    docx_path = markdown_to_docx(markdown, output_dir / f"{safe_section}_report.docx")
    pdf_path = markdown_to_pdf(markdown, output_dir / f"{safe_section}_report.pdf")
    return {"markdown": md_path, "docx": docx_path, "pdf": pdf_path}


def answer_metadata(answer: Any | None) -> dict[str, Any]:
    if answer is None:
        return {}
    if hasattr(answer, "metadata"):
        try:
            return dict(answer.metadata())
        except Exception:  # noqa: BLE001 - report generation should not fail on optional metadata.
            return {}
    if isinstance(answer, dict):
        return answer
    return {
        "provider": getattr(answer, "provider", None),
        "model": getattr(answer, "model", None),
        "status": getattr(answer, "status", None),
    }


def answer_text(answer: Any | None, fallback: str = "") -> str:
    if answer is None:
        return fallback
    return compact_text(getattr(answer, "text", answer), 20_000) or fallback


def _numeric_value(value: Any) -> float | None:
    if value in (None, ""):
        return None
    if isinstance(value, bool):
        return None
    if isinstance(value, (int, float)):
        return float(value)
    text = str(value).strip().replace(",", ".")
    match = re.search(r"-?\d+(?:\.\d+)?", text)
    if not match:
        return None
    try:
        return float(match.group(0))
    except ValueError:
        return None


def _format_numeric(value: float | None) -> int | float | None:
    if value is None:
        return None
    if float(value).is_integer():
        return int(value)
    return round(value, 4)


def _usage_value(usage: dict[str, Any], *keys: str) -> float | None:
    for key in keys:
        if key in usage:
            value = _numeric_value(usage.get(key))
            if value is not None:
                return value
    return None


def _env_float(name: str, default: float) -> float:
    try:
        return float(os.getenv(name, str(default)))
    except (TypeError, ValueError):
        return default


def routerai_budget_summary(answer: Any | None, *, budget_rub: float = DEFAULT_ROUTERAI_BUDGET_RUB) -> dict[str, Any]:
    metadata = answer_metadata(answer)
    usage = metadata.get("usage") if isinstance(metadata.get("usage"), dict) else {}
    prompt_tokens = _usage_value(usage, "prompt_tokens", "input_tokens", "tokens_prompt")
    completion_tokens = _usage_value(usage, "completion_tokens", "output_tokens", "tokens_completion")
    total_tokens = _usage_value(usage, "total_tokens", "tokens_total", "tokens")
    elapsed_seconds = _usage_value(usage, "elapsed_seconds", "duration_seconds", "latency_seconds")
    if total_tokens is None and prompt_tokens is not None and completion_tokens is not None:
        total_tokens = prompt_tokens + completion_tokens

    reported_cost_rub = _usage_value(usage, "cost_rub", "total_cost_rub", "amount_rub", "rub")
    reported_cost = _usage_value(usage, "cost", "total_cost", "amount")
    currency = compact_text(usage.get("currency") or usage.get("cost_currency") or usage.get("billing_currency"), 30)
    prompt_rate = _env_float("ROUTERAI_PROMPT_RUB_PER_1K", DEFAULT_ROUTERAI_PROMPT_RUB_PER_1K)
    completion_rate = _env_float("ROUTERAI_COMPLETION_RUB_PER_1K", DEFAULT_ROUTERAI_COMPLETION_RUB_PER_1K)
    estimated_cost_rub = None
    if reported_cost_rub is None:
        if prompt_tokens is not None or completion_tokens is not None:
            estimated_cost_rub = ((prompt_tokens or 0) / 1000.0 * prompt_rate) + ((completion_tokens or 0) / 1000.0 * completion_rate)
        elif total_tokens is not None:
            estimated_cost_rub = total_tokens / 1000.0 * max(prompt_rate, completion_rate)
    remaining_rub = budget_rub - reported_cost_rub if reported_cost_rub is not None else None
    estimated_remaining_rub = budget_rub - estimated_cost_rub if estimated_cost_rub is not None else None

    if reported_cost_rub is not None and remaining_rub < 0:
        status = "budget_exceeded"
    elif reported_cost_rub is not None:
        status = "cost_recorded"
    elif total_tokens is not None:
        status = "tokens_recorded"
    elif answer is None:
        status = "no_answer"
    else:
        status = "no_usage_metadata"

    note = (
        "Фактическая стоимость в рублях не оценивалась: API не вернул cost_rub. "
        "Контролируем token usage и общий лимит бюджета."
        if reported_cost_rub is None
        else "API вернул стоимость в рублях; остаток рассчитан от демо-лимита."
    )
    return {
        "provider": metadata.get("provider") or "",
        "model": metadata.get("model") or "",
        "status": metadata.get("status") or status,
        "budget_status": status,
        "budget_rub": _format_numeric(float(budget_rub)),
        "prompt_tokens": _format_numeric(prompt_tokens),
        "completion_tokens": _format_numeric(completion_tokens),
        "total_tokens": _format_numeric(total_tokens),
        "reported_cost": _format_numeric(reported_cost),
        "reported_cost_currency": currency,
        "reported_cost_rub": _format_numeric(reported_cost_rub),
        "estimated_cost_rub": _format_numeric(estimated_cost_rub),
        "estimated_remaining_budget_rub": _format_numeric(estimated_remaining_rub),
        "elapsed_seconds": _format_numeric(elapsed_seconds),
        "remaining_budget_rub": _format_numeric(remaining_rub),
        "usage_keys": ", ".join(sorted(str(key) for key in usage.keys())) if usage else "",
        "note": note,
    }


def format_routerai_budget_summary(answer: Any | None, *, budget_rub: float = DEFAULT_ROUTERAI_BUDGET_RUB) -> list[str]:
    summary = routerai_budget_summary(answer, budget_rub=budget_rub)
    lines = [
        f"- Provider/model: {summary.get('provider') or 'n/a'} / {summary.get('model') or 'n/a'}",
        f"- Budget limit: {summary['budget_rub']} RUB",
        f"- Token usage: prompt={summary.get('prompt_tokens') or 'n/a'}, completion={summary.get('completion_tokens') or 'n/a'}, total={summary.get('total_tokens') or 'n/a'}",
    ]
    if summary.get("reported_cost_rub") is not None:
        lines.append(f"- Reported cost: {summary['reported_cost_rub']} RUB; remaining={summary.get('remaining_budget_rub')}")
    elif summary.get("reported_cost") is not None:
        currency = summary.get("reported_cost_currency") or "API currency"
        lines.append(f"- Reported cost: {summary['reported_cost']} {currency}")
    lines.append(f"- Status: {summary['budget_status']}")
    lines.append(f"- Note: {summary['note']}")
    return lines


def answer_report_links(run: Any | None, *, limit: int = 25) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    if run is None:
        return rows
    for index, result in enumerate((getattr(run, "results", []) or [])[:limit], start=1):
        confidence = relevance_confidence(result)
        rows.append(
            {
                "index": index,
                "title": compact_text(getattr(result, "title", ""), 600),
                "year": getattr(result, "year", None),
                "source": getattr(result, "source", ""),
                "score": getattr(result, "score", None),
                "confidence": confidence,
                "link": source_link(result),
            }
        )
    return rows


def clean_report_text(value: Any, max_chars: int | None = None) -> str:
    text = repair_mojibake(value)
    text = XML_INVALID_CHAR_RE.sub(" ", text)
    text = text.replace("**", "")
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.splitlines()]
    cleaned = "\n".join(line for line in lines if line)
    if max_chars is not None and len(cleaned) > max_chars:
        return cleaned[: max_chars - 3].rstrip() + "..."
    return cleaned


def report_text_blocks(text: Any) -> list[dict[str, str]]:
    cleaned = clean_report_text(text, 20_000)
    blocks: list[dict[str, str]] = []
    for raw_line in cleaned.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        heading = ""
        if line.startswith("### "):
            heading = line[4:].strip()
        elif line.startswith("## "):
            heading = line[3:].strip()
        elif line.startswith("# "):
            heading = line[2:].strip()
        else:
            match = re.match(r"^(?:\d+[.)]\s*)?([А-ЯA-ZЁ][^.!?]{2,90})[:：]$", line)
            if match:
                heading = match.group(1).strip()
        if heading:
            blocks.append({"type": "heading", "text": compact_text(heading, 180)})
            continue
        if line.startswith(("- ", "* ", "• ")):
            blocks.append({"type": "bullet", "text": compact_text(line[2:], 1200)})
        else:
            blocks.append({"type": "paragraph", "text": compact_text(line, 1800)})
    return blocks


def answer_report_sections(
    *,
    query: str | None,
    answer: Any | None,
    run: Any | None = None,
    orchestration: Any | None = None,
) -> list[dict[str, Any]]:
    report_query = query or getattr(getattr(run, "request", None), "query", "") or (
        orchestration_query(orchestration) if orchestration is not None else ""
    )
    raw_answer = getattr(answer, "text", None) if answer is not None else None
    if raw_answer in (None, ""):
        raw_answer = "Отчет модели еще не был сгенерирован."
    sections: list[dict[str, Any]] = [
        {
            "title": "Отчет по литературному поиску",
            "paragraphs": [f"Запрос: {compact_text(report_query, 800)}"] if report_query else [],
            "blocks": report_text_blocks(raw_answer),
        }
    ]

    links = answer_report_links(run)
    sections.append(
        {
            "title": "Ключевые web-источники",
            "paragraphs": [] if links else ["Web-источники не найдены."],
            "table": {
                "headers": ["#", "Заголовок", "Ссылка"],
                "rows": [
                    [
                        str(row["index"]),
                        row["title"],
                        row.get("link") or "",
                    ]
                    for row in links[:20]
                ],
            }
            if links
            else None,
        }
    )

    if run is not None:
        local_rows = getattr(run, "local_matches", []) or []
        sections.append(
            {
                "title": "Ключевые локальные источники",
                "paragraphs": [] if local_rows else ["Локальные источники по запросу не найдены."],
                "table": {
                    "headers": ["#", "Источник"],
                    "rows": [
                        [
                            str(index),
                            source_title(row),
                        ]
                        for index, row in enumerate(local_rows[:20], start=1)
                    ],
                }
                if local_rows
                else None,
            }
        )

    if run is not None and getattr(run, "deep_results", None):
        deep_blocks: list[dict[str, str]] = []
        for index, item in enumerate((getattr(run, "deep_results", []) or [])[:20], start=1):
            summary = getattr(item, "document_summary", {}) or {}
            source = getattr(item, "source_result", None)
            title = compact_text(getattr(source, "title", "") if source is not None else "", 260)
            link = source_link(source) if source is not None else ""
            deep_blocks.append({"type": "heading", "text": f"{index}. {title}"})
            if link:
                deep_blocks.append({"type": "paragraph", "text": f"Ссылка: {link}"})
            deep_blocks.append(
                {
                    "type": "paragraph",
                    "text": compact_text(summary.get("summary") or summary.get("main_topic") or "Summary не извлечен.", 1400),
                }
            )
        sections.append({"title": "Deep Search summaries", "paragraphs": [], "blocks": deep_blocks})

    if orchestration is not None:
        rows = orchestration_all_rows(orchestration)
        data_rows = [row for row in rows if row.get("source_type") != "diagnostics"]
        evidence_blocks: list[dict[str, str]] = []
        for index, row in enumerate(data_rows[:25], start=1):
            locator = row_locator(row)
            locator_text = f"; {locator}" if locator else ""
            evidence_blocks.append({"type": "heading", "text": f"{index}. {row_title(row)}{locator_text}"})
            preview = compact_text(row.get("preview") or row.get("summary") or row.get("path"), 900)
            if preview:
                evidence_blocks.append({"type": "paragraph", "text": preview})
        sections.append(
            {
                "title": "Local RAG evidence",
                "paragraphs": [] if evidence_blocks else ["Local RAG evidence не найден."],
                "blocks": evidence_blocks,
            }
        )

    return sections


def build_answer_report_markdown(
    *,
    query: str | None,
    answer: Any | None,
    run: Any | None = None,
    orchestration: Any | None = None,
) -> str:
    lines: list[str] = []
    for section in answer_report_sections(query=query, answer=answer, run=run, orchestration=orchestration):
        lines.extend([f"## {section['title']}", ""])
        for paragraph_text in section.get("paragraphs") or []:
            lines.extend([compact_text(paragraph_text, 1800), ""])
        for block in section.get("blocks") or []:
            if block["type"] == "heading":
                lines.extend([f"### {block['text']}", ""])
            elif block["type"] == "bullet":
                lines.append(f"- {block['text']}")
            else:
                lines.extend([block["text"], ""])
        table = section.get("table")
        if table:
            headers = table["headers"]
            lines.append(" | ".join(headers))
            lines.append(" | ".join("---" for _ in headers))
            for row in table["rows"]:
                lines.append(" | ".join(compact_text(value, 700) for value in row))
            lines.append("")
    return "\n".join(lines).strip() + "\n"


def build_answer_docx_from_sections(sections: list[dict[str, Any]], output_path: Path) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    set_docx_style(document)
    for index, section in enumerate(sections):
        document.add_heading(compact_text(section["title"], 220), level=1 if index == 0 else 2)
        for paragraph_text in section.get("paragraphs") or []:
            document.add_paragraph(compact_text(paragraph_text, 1800))
        for block in section.get("blocks") or []:
            if block["type"] == "heading":
                document.add_heading(compact_text(block["text"], 220), level=3)
            elif block["type"] == "bullet":
                document.add_paragraph(compact_text(block["text"], 1400), style="List Bullet")
            else:
                document.add_paragraph(compact_text(block["text"], 1800))
        table = section.get("table")
        if table:
            add_docx_table(document, table["headers"], table["rows"])
    document.save(output_path)
    return output_path


def build_answer_pdf_from_sections(sections: list[dict[str, Any]], output_path: Path) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    styles = pdf_styles()
    story: list[Any] = []
    for index, section in enumerate(sections):
        story.append(paragraph(section["title"], styles["TitleUnicode"] if index == 0 else styles["HeadingUnicode"]))
        for paragraph_text in section.get("paragraphs") or []:
            story.append(paragraph(paragraph_text, styles["BodyUnicode"]))
        for block in section.get("blocks") or []:
            if block["type"] == "heading":
                story.append(paragraph(block["text"], styles["HeadingUnicode"]))
            elif block["type"] == "bullet":
                story.append(paragraph(f"• {block['text']}", styles["BodyUnicode"]))
            else:
                story.append(paragraph(block["text"], styles["BodyUnicode"]))
        table = section.get("table")
        if table and table.get("rows"):
            story.append(Spacer(1, 0.1 * cm))
            add_pdf_table(story, table["headers"], table["rows"], styles)
        story.append(Spacer(1, 0.25 * cm))
    doc = SimpleDocTemplate(str(output_path), pagesize=A4, rightMargin=1.2 * cm, leftMargin=1.2 * cm, topMargin=1.2 * cm, bottomMargin=1.2 * cm)
    doc.build(story)
    return output_path


def convert_docx_to_pdf(docx_path: Path, output_path: Path) -> Path | None:
    executable = shutil.which("soffice") or shutil.which("libreoffice")
    if not executable:
        return None
    output_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        subprocess.run(
            [executable, "--headless", "--convert-to", "pdf", "--outdir", str(output_path.parent), str(docx_path)],
            check=True,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
            timeout=90,
        )
    except Exception:  # noqa: BLE001 - report export should fall back to the internal PDF renderer.
        return None
    converted = output_path.parent / f"{docx_path.stem}.pdf"
    if converted.exists() and converted != output_path:
        converted.replace(output_path)
    return output_path if output_path.exists() else None


def build_answer_exports(
    output_dir: Path,
    *,
    query: str | None,
    answer: Any | None,
    run: Any | None = None,
    orchestration: Any | None = None,
    prefix: str = "routerai_answer",
) -> dict[str, Path]:
    output_dir.mkdir(parents=True, exist_ok=True)
    safe_prefix = re.sub(r"[^A-Za-z0-9_.-]+", "_", prefix.lower()).strip("_") or "routerai_answer"
    sections = answer_report_sections(query=query, answer=answer, run=run, orchestration=orchestration)
    markdown = build_answer_report_markdown(query=query, answer=answer, run=run, orchestration=orchestration)
    md_path = write_text_report(output_dir / f"{safe_prefix}.md", markdown)
    docx_path = build_answer_docx_from_sections(sections, output_dir / f"{safe_prefix}.docx")
    pdf_path = convert_docx_to_pdf(docx_path, output_dir / f"{safe_prefix}.pdf") or build_answer_pdf_from_sections(
        sections,
        output_dir / f"{safe_prefix}.pdf",
    )
    json_path = write_json_report(
        output_dir / f"{safe_prefix}.json",
        {
            "query": query,
            "answer": answer_text(answer),
            "metadata": answer_metadata(answer),
            "routerai_budget": routerai_budget_summary(answer),
            "web_links": answer_report_links(run),
        },
    )
    return {"markdown": md_path, "pdf": pdf_path, "docx": docx_path, "json": json_path}


def orchestration_query(orchestration: Any, query: str | None = None) -> str:
    rewrite = getattr(orchestration, "query_rewrite", None) or {}
    plan = getattr(orchestration, "plan", None)
    return compact_text(
        query
        or rewrite.get("corrected_query")
        or getattr(plan, "original_query", "")
        or "",
        500,
    )


def orchestration_context(orchestration: Any) -> dict[str, list[dict[str, Any]]]:
    context = getattr(orchestration, "retrieved_context", None)
    if context is None:
        return {"raw": [], "summaries": [], "tables": [], "graph": [], "web": []}
    if hasattr(context, "as_dict"):
        return context.as_dict()
    return dict(context)


def orchestration_all_rows(orchestration: Any) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for section, section_rows in orchestration_context(orchestration).items():
        for row in section_rows or []:
            enriched = dict(row)
            enriched.setdefault("route", section)
            rows.append(enriched)
    return rows


def row_locator(row: dict[str, Any]) -> str:
    return compact_text(
        row.get("url")
        or row.get("doi")
        or row.get("local_path")
        or row.get("source_path")
        or row.get("path")
        or row.get("node_id")
        or "",
        500,
    )


def row_title(row: dict[str, Any]) -> str:
    return compact_text(
        row.get("title")
        or row.get("label")
        or row.get("doc_id")
        or row.get("id")
        or row.get("source_path")
        or "Evidence",
        260,
    )


def append_orchestration_rows(lines: list[str], title: str, rows: list[dict[str, Any]], *, limit: int = 20) -> None:
    lines.extend(["", f"## {title}"])
    data_rows = [row for row in rows if row.get("source_type") != "diagnostics"]
    if not data_rows:
        lines.append("- Нет данных.")
        return
    for index, row in enumerate(data_rows[:limit], start=1):
        score = row.get("score")
        score_text = f"; score={score}" if score not in (None, "") else ""
        locator = row_locator(row)
        locator_text = f"; {locator}" if locator else ""
        preview = compact_text(row.get("preview") or row.get("summary") or row.get("path") or row.get("relation"), 500)
        lines.append(f"{index}. {row_title(row)}{score_text}{locator_text}")
        if preview:
            lines.append(f"   - {preview}")


def orchestration_section_markdown(
    orchestration: Any,
    section: str,
    *,
    answer: Any | None = None,
    query: str | None = None,
) -> str:
    section = section.lower().strip()
    context = orchestration_context(orchestration)
    plan = getattr(orchestration, "plan", None)
    rewrite = getattr(orchestration, "query_rewrite", None) or {}
    lines = [
        f"# RAG отчет: {orchestration_query(orchestration, query)}",
        "",
        f"- Intent: {getattr(plan, 'intent', 'n/a')}",
        f"- Routes: {', '.join(getattr(plan, 'routes', []) or []) or 'n/a'}",
        f"- Answer format: {getattr(plan, 'answer_format', 'n/a')}",
    ]
    if rewrite:
        lines.append(f"- Rewrite: {compact_text(rewrite.get('corrected_query'), 500)}")
        lines.append(f"- Rewrite LLM: {rewrite.get('rewrite_used_llm')}")

    if section in {"answer", "full"}:
        lines.extend(["", "## Ответ", answer_text(answer, getattr(orchestration, "answer_draft", "")) or "Ответ не сгенерирован."])
        metadata = answer_metadata(answer)
        if metadata:
            lines.extend(["", "## LLM metadata"])
            for key, value in metadata.items():
                lines.append(f"- {key}: {compact_text(value, 500)}")

    if section in {"sources", "evidence", "full"}:
        append_orchestration_rows(lines, "Raw RAG", context.get("raw") or [])
        append_orchestration_rows(lines, "Summary RAG", context.get("summaries") or [])
        append_orchestration_rows(lines, "Tables", context.get("tables") or [])
        append_orchestration_rows(lines, "Graph", context.get("graph") or [])
        append_orchestration_rows(lines, "Web", context.get("web") or [])

    if section in {"comparison", "full"}:
        lines.extend(["", "## Local vs Web / Method comparison"])
        local_count = sum(len(context.get(name) or []) for name in ("raw", "summaries", "tables", "graph"))
        web_count = len(context.get("web") or [])
        lines.append(f"- Local evidence rows: {local_count}")
        lines.append(f"- Web evidence rows: {web_count}")
        procedure_rows = [
            row for row in [*(context.get("summaries") or []), *(context.get("web") or [])]
            if "procedure" in compact_text(row.get("source_type") or row.get("kind") or row.get("id")).casefold()
        ]
        append_orchestration_rows(lines, "Методики и режимы", procedure_rows, limit=30)

    if section in {"properties", "property", "numeric", "full"}:
        property_rows = property_report_rows_from_orchestration(orchestration)
        lines.extend(["", property_report_markdown("Свойства и численные результаты", orchestration_query(orchestration, query), property_rows).strip()])

    if section in {"graphs", "graph", "full"}:
        append_orchestration_rows(lines, "Graph evidence", context.get("graph") or [], limit=40)

    if section in {"charts", "full"}:
        lines.extend(["", "## Coverage"])
        for name, rows in context.items():
            data_count = len([row for row in rows if row.get("source_type") != "diagnostics"])
            lines.append(f"- {name}: {data_count}")

    fallbacks = getattr(orchestration, "fallbacks", []) or []
    if fallbacks and section in {"evidence", "full"}:
        lines.extend(["", "## Fallbacks"])
        for item in fallbacks:
            lines.append(f"- {compact_text(item, 800)}")
    return "\n".join(lines).strip() + "\n"


def build_orchestration_exports(
    orchestration: Any,
    section: str,
    output_dir: Path | None = None,
    *,
    answer: Any | None = None,
    query: str | None = None,
) -> dict[str, Path]:
    output_dir = output_dir or PROJECT_ROOT / "data" / "processed" / "rag_runs" / safe_report_id(orchestration_query(orchestration, query))
    output_dir.mkdir(parents=True, exist_ok=True)
    safe_section = re.sub(r"[^A-Za-z0-9_.-]+", "_", section.lower()).strip("_") or "section"
    markdown = orchestration_section_markdown(orchestration, safe_section, answer=answer, query=query)
    md_path = write_text_report(output_dir / f"{safe_section}_report.md", markdown)
    docx_path = markdown_to_docx(markdown, output_dir / f"{safe_section}_report.docx")
    pdf_path = markdown_to_pdf(markdown, output_dir / f"{safe_section}_report.pdf")
    return {"markdown": md_path, "docx": docx_path, "pdf": pdf_path}


def write_orchestration_payload(orchestration: Any, output_path: Path, *, answer: Any | None = None) -> Path:
    payload = orchestration.as_dict() if hasattr(orchestration, "as_dict") else dict(orchestration)
    if answer is not None:
        payload = {**payload, "answer": {"text": answer_text(answer), "metadata": answer_metadata(answer)}}
    return write_json_report(output_path, payload)


def write_orchestration_web_manifest(orchestration: Any, output_path: Path) -> Path:
    rows = []
    for index, row in enumerate(orchestration_context(orchestration).get("web") or [], start=1):
        rows.append(
            {
                "index": index,
                "title": row_title(row),
                "source": row.get("source"),
                "doi": row.get("doi"),
                "url": row.get("url") or row.get("doi"),
                "score": row.get("score"),
                "keyword_hits": row.get("keyword_hits") or [],
            }
        )
    return write_json_report(output_path, {"web_links": rows, "note": "Web full text is not archived; links and summaries are included when available."})


def write_orchestration_local_manifest(
    orchestration: Any,
    output_path: Path,
    *,
    project_root: Path | None = None,
) -> tuple[Path, list[dict[str, Any]]]:
    rows = [
        row
        for row in orchestration_all_rows(orchestration)
        if any(row.get(key) for key in ("local_path", "source_path", "path", "file_name"))
    ]
    return write_local_files_manifest(SimpleNamespace(local_matches=rows), output_path, project_root=project_root)


def write_links_csv(run: Any, output_path: Path) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    with output_path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["#", "title", "year", "source", "confidence", "confidence_percent", "score", "link"])
        for index, result in enumerate(getattr(run, "results", []) or [], start=1):
            confidence = relevance_confidence(result)
            writer.writerow(
                [
                    index,
                    compact_text(result.title, 500),
                    result.year or "",
                    result.source,
                    confidence["label"],
                    confidence["confidence"],
                    confidence["score"],
                    source_link(result),
                ]
            )
    return output_path


def write_web_links_manifest(run: Any, output_path: Path) -> Path:
    rows = []
    for index, result in enumerate(getattr(run, "results", []) or [], start=1):
        rows.append(
            {
                "index": index,
                "title": compact_text(result.title, 600),
                "year": result.year,
                "source": result.source,
                "doi": result.doi,
                "url": source_link(result),
                "open_access": getattr(result, "open_access", {}) or {},
                "score": result.score,
                "relevance_confidence": relevance_confidence(result),
                "keyword_hits": result.keyword_hits,
            }
        )
    return write_json_report(output_path, {"web_links": rows, "note": "Full web texts are not archived; links and extracted summaries are included."})


def local_candidate_values(row: dict[str, Any]) -> list[str]:
    values: list[str] = []
    for key in ("local_path", "source_path", "path", "file_name"):
        value = row.get(key)
        if value:
            values.append(str(value))
    return values


def is_within(path: Path, root: Path) -> bool:
    try:
        path.resolve().relative_to(root.resolve())
        return True
    except ValueError:
        return False


def find_local_file(value: str, *, project_root: Path | None = None) -> Path | None:
    project_root = project_root or PROJECT_ROOT
    if not value:
        return None
    candidate = Path(value)
    if candidate.is_absolute() and candidate.exists() and candidate.is_file() and is_within(candidate, project_root):
        return candidate
    if not candidate.is_absolute():
        joined = project_root / candidate
        if joined.exists() and joined.is_file() and is_within(joined, project_root):
            return joined
    file_name = Path(value).name
    if not file_name:
        return None
    raw_root = project_root / "data" / "raw"
    if raw_root.exists():
        for match in raw_root.rglob(file_name):
            if match.is_file() and is_within(match, project_root):
                return match
    return None


def write_local_files_manifest(run: Any, output_path: Path, *, project_root: Path | None = None) -> tuple[Path, list[dict[str, Any]]]:
    project_root = project_root or PROJECT_ROOT
    manifest: list[dict[str, Any]] = []
    seen: set[Path] = set()
    total_bytes = 0
    included_count = 0
    for index, row in enumerate(getattr(run, "local_matches", []) or [], start=1):
        found: Path | None = None
        for value in local_candidate_values(row):
            found = find_local_file(value, project_root=project_root)
            if found:
                break
        item = {
            "index": index,
            "title": source_title(row),
            "doc_id": row.get("doc_id"),
            "source_path": row.get("source_path"),
            "local_path": str(found) if found else row.get("local_path"),
            "archive_path": None,
            "status": "missing_local_file",
        }
        if found:
            resolved = found.resolve()
            size = found.stat().st_size
            if resolved in seen:
                item["status"] = "duplicate"
            elif included_count >= MAX_LOCAL_ARCHIVE_FILES:
                item["status"] = "skipped_file_limit"
            elif total_bytes + size > MAX_LOCAL_ARCHIVE_BYTES:
                item["status"] = "skipped_size_limit"
            else:
                seen.add(resolved)
                included_count += 1
                total_bytes += size
                item["status"] = "included"
                item["archive_path"] = f"local_publications/{included_count:02d}_{found.name}"
                item["size_bytes"] = size
        manifest.append(item)
    return write_json_report(
        output_path,
        {
            "limits": {"max_files": MAX_LOCAL_ARCHIVE_FILES, "max_bytes": MAX_LOCAL_ARCHIVE_BYTES},
            "included_files": included_count,
            "included_bytes": total_bytes,
            "local_files": manifest,
        },
    ), manifest


def archive_safe_name(value: Any, *, fallback: str = "source", max_chars: int = 90) -> str:
    text = compact_text(value, max_chars=max_chars) or fallback
    text = XML_INVALID_CHAR_RE.sub(" ", text)
    text = re.sub(r'[<>:"/\\|?*\x00-\x1F]+', "_", text)
    text = re.sub(r"\s+", " ", text).strip(" ._")
    return text[:max_chars].strip(" ._") or fallback


def unique_archive_name(name: str, used: set[str]) -> str:
    candidate = name
    stem = Path(name).stem
    suffix = Path(name).suffix
    counter = 2
    while candidate.casefold() in used:
        candidate = f"{stem}_{counter}{suffix}"
        counter += 1
    used.add(candidate.casefold())
    return candidate


def build_local_publications_archive(
    run: Any,
    output_path: Path,
    *,
    project_root: Path | None = None,
) -> Path:
    project_root = project_root or PROJECT_ROOT
    output_path.parent.mkdir(parents=True, exist_ok=True)
    manifest_path, manifest = write_local_files_manifest(
        run,
        output_path.with_suffix(".manifest.json"),
        project_root=project_root,
    )
    del manifest_path
    included = [item for item in manifest if item.get("status") == "included" and item.get("local_path")]
    used_names: set[str] = set()
    with zipfile.ZipFile(output_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        toc = "rank;title;file_name;status\n"
        for item in manifest:
            toc += (
                f"{item.get('index')};"
                f"{compact_text(item.get('title'), 400).replace(';', ',')};"
                f"{Path(str(item.get('local_path') or '')).name};"
                f"{item.get('status')}\n"
            )
        zf.writestr("sources.csv", toc.encode("utf-8-sig"))
        for item in included:
            source_path = Path(str(item["local_path"]))
            if not source_path.exists() or not source_path.is_file() or not is_within(source_path, project_root):
                continue
            title = archive_safe_name(item.get("title") or source_path.stem, fallback=source_path.stem)
            arcname = unique_archive_name(f"{int(item.get('index') or 0):02d}_{title}{source_path.suffix}", used_names)
            zf.write(source_path, arcname=arcname)
        if not included:
            zf.writestr(
                "README.txt",
                "Локальные файлы для найденных публикаций не обнаружены. Проверьте sources.csv.\n".encode("utf-8"),
            )
    return output_path


def build_web_publications_archive(run: Any, output_path: Path) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    used_names: set[str] = set()
    with zipfile.ZipFile(output_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        toc = "rank;title;year;source;link\n"
        results = list(getattr(run, "results", []) or [])
        for index, result in enumerate(results, start=1):
            title = compact_text(getattr(result, "title", ""), 500)
            link = preferred_web_link(result)
            toc += (
                f"{index};"
                f"{title.replace(';', ',')};"
                f"{getattr(result, 'year', '') or ''};"
                f"{getattr(result, 'source', '') or ''};"
                f"{link}\n"
            )
            if not link:
                continue
            shortcut_name = unique_archive_name(
                f"{index:02d}_{archive_safe_name(title, fallback='web_source')}.url",
                used_names,
            )
            zf.writestr(
                shortcut_name,
                f"[InternetShortcut]\r\nURL={link}\r\n".encode("utf-8"),
            )
        zf.writestr("sources.csv", toc.encode("utf-8-sig"))
        if not results:
            zf.writestr("README.txt", "Web-источники не найдены.\n".encode("utf-8"))
    return output_path


def build_run_archive(
    run: Any,
    output_path: Path,
    *,
    project_root: Path | None = None,
    answer: Any | None = None,
    query: str | None = None,
) -> Path:
    project_root = project_root or PROJECT_ROOT
    output_path.parent.mkdir(parents=True, exist_ok=True)
    run_dir = Path(getattr(run, "output_dir", "") or output_path.parent)
    links_csv = write_links_csv(run, run_dir / "links.csv")
    web_manifest = write_web_links_manifest(run, run_dir / "web_links_manifest.json")
    local_manifest_path, local_manifest = write_local_files_manifest(run, run_dir / "local_publication_files_manifest.json", project_root=project_root)
    candidate_paths: list[Path] = [
        run_dir / "request.json",
        run_dir / "query_plan.json",
        run_dir / "keywords.json",
        run_dir / "metadata_results.jsonl",
        run_dir / "local_matches.jsonl",
        run_dir / "resource_links.jsonl",
        run_dir / "comparison_report.json",
        run_dir / "literature_report.md",
        run_dir / "literature_links_report.md",
        run_dir / "deep_search_report.md",
        run_dir / "executive_brief.md",
        run_dir / "full_run.json",
        run_dir / "web_document_summaries.jsonl",
        run_dir / "web_procedure_summaries.jsonl",
        run_dir / "deep_search_results.jsonl",
        links_csv,
        web_manifest,
        local_manifest_path,
    ]
    for section in ("sources", "comparison", "properties", "evidence", "graphs", "charts", "deep"):
        candidate_paths.extend(build_section_exports(run, section, run_dir / "section_reports").values())
    if answer is not None:
        candidate_paths.extend(
            build_answer_exports(
                run_dir / "answer_report",
                query=query or getattr(getattr(run, "request", None), "query", None),
                answer=answer,
                run=run,
            ).values()
        )
    for attr in (
        "report_pdf_path",
        "report_docx_path",
        "links_report_pdf_path",
        "links_report_docx_path",
        "deep_report_pdf_path",
        "deep_report_docx_path",
        "executive_brief_pdf_path",
        "executive_brief_docx_path",
    ):
        value = getattr(run, attr, None)
        if value:
            candidate_paths.append(Path(value))
    seen: set[Path] = set()
    with zipfile.ZipFile(output_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for path in candidate_paths:
            if not path or not path.exists() or path in seen:
                continue
            seen.add(path)
            try:
                arcname = path.relative_to(run_dir)
            except ValueError:
                arcname = Path(path.name)
            zf.write(path, arcname=str(arcname))
        for item in local_manifest:
            if item.get("status") != "included" or not item.get("archive_path") or not item.get("local_path"):
                continue
            source_path = Path(str(item["local_path"]))
            if source_path.exists() and source_path.is_file() and is_within(source_path, project_root):
                zf.write(source_path, arcname=str(item["archive_path"]))
    return output_path


def build_orchestration_archive(
    orchestration: Any,
    output_path: Path,
    *,
    answer: Any | None = None,
    query: str | None = None,
    project_root: Path | None = None,
) -> Path:
    project_root = project_root or PROJECT_ROOT
    output_path.parent.mkdir(parents=True, exist_ok=True)
    run_dir = output_path.parent
    payload_path = write_orchestration_payload(orchestration, run_dir / "orchestration_payload.json", answer=answer)
    web_manifest = write_orchestration_web_manifest(orchestration, run_dir / "orchestration_web_links_manifest.json")
    local_manifest_path, local_manifest = write_orchestration_local_manifest(
        orchestration,
        run_dir / "orchestration_local_files_manifest.json",
        project_root=project_root,
    )
    candidate_paths: list[Path] = [
        payload_path,
        web_manifest,
        local_manifest_path,
    ]
    for section in ("full", "answer", "sources", "comparison", "properties", "evidence", "graphs", "charts"):
        candidate_paths.extend(
            build_orchestration_exports(
                orchestration,
                section,
                run_dir / "section_reports",
                answer=answer,
                query=query,
            ).values()
        )

    seen: set[Path] = set()
    with zipfile.ZipFile(output_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for path in candidate_paths:
            if not path or not path.exists() or path in seen:
                continue
            seen.add(path)
            try:
                arcname = path.relative_to(run_dir)
            except ValueError:
                arcname = Path(path.name)
            zf.write(path, arcname=str(arcname))
        for item in local_manifest:
            if item.get("status") != "included" or not item.get("archive_path") or not item.get("local_path"):
                continue
            source_path = Path(str(item["local_path"]))
            if source_path.exists() and source_path.is_file() and is_within(source_path, project_root):
                zf.write(source_path, arcname=str(item["archive_path"]))
    return output_path
